import pdf2docx
from docx import Document
import os 
import re
from datetime import datetime

# TODO : Error with reading data from June - better parser for hours, takes into account rows
# TODO : update README


def parse_workday_docx(document: Document):
    ''' helper that gets raw data from workday docx '''
    table = document.tables[1]
    data = []
    keys = None
    # every second table out of 3 is the table that we want
    for i in range(1, len(document.tables), 3):
        for i, row in enumerate(document.tables[i].rows):
            text = [cell.text for cell in row.cells]
            if i == 0:
                continue
            if i == 1:
                if keys == None:
                    keys = tuple(text)
            else:
                row_data = dict(zip(keys, text))
                data.append(row_data)
    return data


def fetch_workday_data(pdf_path):
    '''
    creates empty docx, convert workday output to docx and saves,
    then gets raw data from new docx with above helper
    '''
    docx_path = "./script/blank.docx"
    if not os.path.exists(docx_path):
        document = Document()
        document.save(docx_path)
        os.chmod(docx_path, 0o775)

    pdf2docx.parse(pdf_path, docx_path)
    if not os.path.exists(docx_path):
        exit(f"Failed to create file {docx_path}.")
    document = Document(docx_path)

    raw_data = parse_workday_docx(document)
    name, start_date, end_date = document.paragraphs[1].text.split('\n')
    start_date = start_date[-10:].strip()
    end_date = end_date[-10:].strip()
    name = name.strip()
    return raw_data, name, start_date, end_date


def organize_data(raw_data):
    '''
    returns data in format [comment, date, hours logged]
    combines entries that are logged on the same day
    '''
    entries = []
    temp = []
    r = re.compile('[0-9]*/[0-9]*/[0-9]*')
    h = re.compile('Hours:.*')
    for _, entry in enumerate(raw_data):     
        if r.match(entry['Date']) is not None:
            # date for the current entry
            if len(temp) >= 2:
                continue
            temp.append(entry['Date'])
        elif h.match(entry['Date']) is not None:
            # hours for the current entry
            hours = re.findall("\d+\.?\d*", entry['Date'])
            if hours:
                temp.append(hours[0])
                entries.append(temp)
                # reset temp
                temp = []
        else:
            # comment for the current entry
            if len(temp) >= 2:
                temp[0] += ', ' + entry['Comment']
            else:
                temp.append(entry['Comment'])
    return entries


def fill_document(data, name, start_date, end_date):
    PAR_path = "./par-template/PAR-template.docx"
    PAR = Document(PAR_path)

    # fill in header info
    PAR.tables[2].rows[0].cells[1].text.replace("XX", name)
    PAR.tables[2].rows[2].cells[3].text = start_date
    PAR.tables[2].rows[2].cells[5].text = end_date

    total_hours = 0
    # sublist: 0=description, 1=date, 2=hours, PAR table: 0=date, 1=description, 2=hours
    for i, sublist in enumerate(data):
        if not sublist[0]:
            print(f"NOTE: {name} didn't include a description for all days worked.")
        PAR.tables[3].rows[i+1].cells[1].text = sublist[0]
        PAR.tables[3].rows[i+1].cells[0].text = sublist[1]
        PAR.tables[3].rows[i+1].cells[2].text = sublist[2]
        if sublist[2]:
            total_hours += float(sublist[2])
    PAR.tables[3].rows[-1].cells[2].text = str(total_hours)
    
    name = name.replace(" ", "-")
    start_date = start_date.replace("/", ".")
    end_date = end_date.replace("/", ".")
    filename = f'./filled-reports/{name}_{start_date}_to_{end_date}.docx'
    open(filename, "w").close()
    PAR.save(filename)
    os.chmod(filename, 0o666)  # gives permissions rw-rw-rw-


def compare_dates(current_start, current_end, new_start, new_end):
    '''
    compares current start and end dates with new start and end dates 
    returns updated start and end dates 
    '''
    if not current_start:
        current_start = new_start
    else:
        current_start = min(current_start, new_start, key=lambda date: datetime.strptime(date, "%m/%d/%Y"))
    
    if not current_end:
        current_end = new_end
    else:
        current_end = max(current_end, new_end, key=lambda date: datetime.strptime(date, "%m/%d/%Y"))
    
    return current_start, current_end


def main():
    employees = os.listdir("./employees")
    # loop through employees
    for _, employee in enumerate(employees):
        # resetting initial variables
        organized_data = []
        start_date, end_date = "", ""
        employee_name = ""
        employee_path = f"./employees/{employee}"
        print(f"Processing {employee_path}...")
        for file in os.listdir(employee_path):
            raw_data, name, file_start, file_end = fetch_workday_data(f"{employee_path}/{file}")
            # update values
            start_date, end_date = compare_dates(start_date, end_date, file_start, file_end)
            employee_name = name
            if raw_data:
                # add organized data to list of all data
                organized_data += organize_data(raw_data)
            else:
                exit("Fetching workday data failed. Exiting")
        # fill document for current employee
        fill_document(organized_data, employee_name, start_date, end_date)
        print(f"Successfully created report for {name}")


if __name__ == "__main__":
    main()