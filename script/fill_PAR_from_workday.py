from docx import Document
from docx.shared import Pt
from pdf2docx import parse
import os 
import re



def parse_workday_docx(document: Document):
    table = document.tables[1]
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = [cell.text for cell in row.cells]
        if i == 0:
            continue
        if i == 1:
            keys = tuple(text)
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return data


def fetch_workday_data(path):
    # create docx in directory (proxy for saving pdf as docx)
    docx_path = "./unfilled-reports/test.docx"
    document = Document()
    os.chmod(docx_path, 0o775)
    document.save(f"{docx_path}")

    # convert existing pdf to docx for parsing
    pdf_path = f"./employee-pdfs/{path}"

    try:
        parse(pdf_path, docx_path)
    except FileNotFoundError as fnf_error:
        print(f"File Not Found : {fnf_error}")
    except Exception as e:
        print(f'Error converting pdf to docx : {e}')

    # save and parse dat
    document = Document(docx_path)
    raw_data = parse_workday_docx(document)

    # get employee info
    name, start_date, end_date = document.paragraphs[1].text.split('\n')
    start_date = start_date[-10:].strip()
    end_date = end_date[-10:].strip()
    name = name.strip()

    # organize dat
    return raw_data, name, start_date, end_date


def organize_data(raw_data):
    '''
    returns data in format : [comment, date, hours logged]
    combines entries that are logged on the same day
    '''
    entries = []
    temp = []
    r = re.compile('[0-9]*/[0-9]*/[0-9]*')
    h = re.compile('Hours:.*')
    for index, entry in enumerate(raw_data):
        if index == 0:
            continue
        else:
            if r.match(entry['Date']) is not None:
                # date for the current entry
                if len(temp) >= 2:
                    continue
                temp.append(entry['Date'])
            elif h.match(entry['Date']) is not None:
                # hours for the current entry
                temp.append(entry['Date'][9:])
                entries.append(temp)
                # reset temp
                temp = []
            else:
                # comment for the current entry
                description = entry['Comment'].replace('\n', '')
                if len(temp) >= 2:
                    temp[0] += ';\n' + description
                else:
                    temp.append(description)
    return entries

# NOTE FORGERY CONCERNS WITH SIGNATURE, DISCUSS LATER
def fill_document(data, name, start_date, end_date):
    PAR_path = "./par-template/PAR-template.docx"
    PAR = Document(PAR_path)

    # name
    PAR.tables[2].rows[0].cells[1].text.replace("XX", name)
    # period beginning
    PAR.tables[2].rows[2].cells[3].text = start_date
    # period ending
    PAR.tables[2].rows[2].cells[5].text = end_date

    # PAR Template is 1 page with 20 rows, which is enough because there should be a max of 14 rows for each 2 week period
    if len(data) > 20:
        exit("Your Workday pdf output has too many time entries (over 20) for the report. Please adjust and resubmit. Exiting")
    
    # sublist indices: 0 = description, 1 = date, 2 = hours worked
    # desired format : 0 = date, 1 = description, 2 = hours worked
    total_hours_worked = 0
    for i, sublist in enumerate(data):
        for j, item in enumerate(sublist):
            # description
            if j == 0:
                PAR.tables[3].rows[i+1].cells[1].text = item
            # date
            elif j == 1:
                PAR.tables[3].rows[i+1].cells[0].text = item
            # hours worked
            else:
                PAR.tables[3].rows[i+1].cells[j].text = item
                total_hours_worked += float(item)
    
    # add total hours worked to the end of sheet
    PAR.tables[3].rows[-1].cells[-1].text = str(total_hours_worked)

    # TODO : auto fill in date/signature?
    
    # '/' and ' ' interfere with the naming of files
    start_date = start_date.replace("/", "_")
    name = name.replace(" ", "")
    filename = f'./filled-reports/{name}/PAR-{name}-{start_date}.docx'

    # check if directory exists
    if not os.path.exists(f'./filled-reports/{name}'):
        os.makedirs(f'filled-reports/{name}')
    
    # save file with filename at directory
    try :
        PAR.save(filename)
    except Exception as e:
        print(f"Error saving document : {e}")
        raise

    print(f"File for {name} successfully generated for {start_date}")
    # return path to successfully generated file
    return filename

def combine_documents(filenames, filename):
    # use first doc as default template so styling is consistent
    merged_document = Document(filenames[0])
    for index, file in enumerate(filenames):
        if index == 0:
            merged_document.add_page_break()
            continue
        else:
            cur_doc = Document(file)
            
            if index != len(filenames) - 1:
                cur_doc.add_page_break()
            
            for element in cur_doc.element.body:
                merged_document.element.body.append(element)
    merged_document.save(filename)
    return filename
    

def main():
    # assuming that each employee has their own directory with their pdfs in it in employee-pdfs
    for dir in os.listdir("./employee-pdfs"):
        if os.path.isfile("dir"):
            print(f"file object found at {dir}")
            continue
        else:
            # list of successfully created file paths for the specific employee
            file_paths = []

            # creating each file individually
            for file in os.listdir(f'./employee-pdfs/{dir}'):
                raw_data, name, start_date, end_date = fetch_workday_data(f"{dir}/{file}")
                if raw_data:
                    organized_data = organize_data(raw_data)
                    path = fill_document(organized_data, name, start_date, end_date)
                    file_paths.append(path)
                else:
                    exit(1, f"Fetching workday data for ./employee-pdfs/{dir} failed. Exiting")
            
            # adds all of the docs together
            combine_documents(file_paths, f"./final-reports/final-report-{dir}.docx")


if __name__ == "__main__":
    main()